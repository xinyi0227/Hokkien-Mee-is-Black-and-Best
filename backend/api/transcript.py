import os
import azure.cognitiveservices.speech as speechsdk
from django.conf import settings
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from .models import Meeting, MeetingFile, Employee, Department,Task, Complaint
from docx import Document
import re
from .views import get_meeting_summary_and_tasks
from django.utils import timezone 
from django.core.files.base import ContentFile
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import io
import json
from datetime import datetime
from django.http import JsonResponse
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

def azure_transcribe(file_path):
    speech_config = speechsdk.SpeechConfig(
        subscription=os.getenv("AZURE_KEY"),
        region=os.getenv("AZURE_REGION")
    )
    audio_config = speechsdk.audio.AudioConfig(filename=file_path)
    speech_recognizer = speechsdk.SpeechRecognizer(
        speech_config=speech_config,
        audio_config=audio_config
    )

    transcript_parts = []

    def handle_result(evt):
        if evt.result.reason == speechsdk.ResultReason.RecognizedSpeech:
            transcript_parts.append(evt.result.text)

    # Subscribe to events
    speech_recognizer.recognized.connect(handle_result)

    # Start recognition
    speech_recognizer.start_continuous_recognition()
    speech_recognizer.session_stopped.connect(lambda evt: speech_recognizer.stop_continuous_recognition())
    speech_recognizer.canceled.connect(lambda evt: speech_recognizer.stop_continuous_recognition())

    # Wait until recognition completes
    done = False
    def stop_cb(evt):
        nonlocal done
        done = True
    speech_recognizer.session_stopped.connect(stop_cb)
    speech_recognizer.canceled.connect(stop_cb)

    while not done:
        import time
        time.sleep(0.5)

    return " ".join(transcript_parts)



@csrf_exempt
def transcript_view(request, meeting_id):
    if request.method != "POST":
        return JsonResponse({"error": "Invalid method"}, status=405)

    try:
        # 1️⃣ Get meeting
        meeting = Meeting.objects.get(meeting_id=meeting_id)

        # --- Resolve mic employees ---
        mic_employees = []
        for mic_field in [meeting.meeting_mic1, meeting.meeting_mic2, meeting.meeting_mic3]:
            if mic_field:
                try:
                    emp = Employee.objects.get(employee_id=mic_field)
                    mic_employees.append(emp.employee_name)
                except Employee.DoesNotExist:
                    mic_employees.append(None)
            else:
                mic_employees.append(None)

        # --- Resolve departments ---
        dept_names = []
        if meeting.meeting_department:
            dept_ids = meeting.meeting_department.split(",")
            departments = Department.objects.filter(department_id__in=dept_ids)
            dept_names = [d.department_name for d in departments]

        # --- Resolve participants ---
        participant_names = []
        if meeting.meeting_participant:
            participant_ids = meeting.meeting_participant.split(",")
            participants = Employee.objects.filter(employee_id__in=participant_ids)
            participant_names = [p.employee_name for p in participants]

        meeting_data = {
            "ID": meeting.meeting_id,
            "title": meeting.meeting_title,
            "date": meeting.meeting_date,
            "time": meeting.meeting_time,
            "location": meeting.meeting_location,
            "mics": mic_employees,
            "departments": dept_names,
            "participants": participant_names
        }

        # 2️⃣ Get meeting files
        meeting_files = MeetingFile.objects.filter(meeting_id=meeting_id)
        file_urls = []
        transcript_file_urls = []

        safe_title = re.sub(r'[^A-Za-z0-9_-]+', '_', meeting.meeting_title)

        for mf in meeting_files:
            if mf.meeting_org:
                file_path = os.path.join(settings.MEDIA_ROOT, mf.meeting_org.name)
                transcript_text = azure_transcribe(file_path)

                # 📂 Make transcripts folder
                transcript_dir = os.path.join(settings.MEDIA_ROOT, "transcripts")
                os.makedirs(transcript_dir, exist_ok=True)

                # Save transcript to TXT
                txt_filename = f"transcript_meeting_{safe_title}.txt"
                txt_path = os.path.join(transcript_dir, txt_filename)
                with open(txt_path, "w", encoding="utf-8") as f:
                    f.write(transcript_text)

                # Save transcript to Word
                doc_filename = f"transcript_meeting_{safe_title}.docx"
                doc_path = os.path.join(transcript_dir, doc_filename)
                document = Document()
                document.add_heading(f"Transcript for Meeting {meeting.meeting_title}", level=1)
                document.add_paragraph(transcript_text)
                document.save(doc_path)

                # ✅ Save only filename (not path) into DB
                mf.meeting_transcripts = doc_filename
                mf.save()

                # Add URLs for frontend download
                txt_url = request.build_absolute_uri(settings.MEDIA_URL + f"transcripts/{txt_filename}")
                doc_url = request.build_absolute_uri(settings.MEDIA_URL + f"transcripts/{doc_filename}")
                transcript_file_urls.extend([txt_url, doc_url])

                # Also return audio file URL
                audio_url = request.build_absolute_uri(settings.MEDIA_URL + mf.meeting_org.name)
                file_urls.append(audio_url)

                # After transcription and saving
                gemini_result = get_meeting_summary_and_tasks(
                    meeting_data,
                    transcript_text,
                    transcript_file_urls
                )

                # Save tasks to DB
                if "tasks" in gemini_result:
                    for name, tasks in gemini_result["tasks"].items():
                        try:
                            assignee = Employee.objects.get(employee_name=name)
                        except Employee.DoesNotExist:
                            assignee = None  # fallback if no match

                        # for task in tasks:
                        #     Task.objects.create(
                        #         meeting=meeting,
                        #         assignee=assignee,
                        #         task_title=task.get("task_title", "Untitled Task"),
                        #         task_content=task.get("task_content", ""),
                        #         urgent_level=task.get("urgent_level", "medium"),
                        #         deadline=task.get("deadline") if task.get("deadline") != "None" else None,
                        #         status="pending",
                        #         created_at=timezone.now()
                        #     )


            # Handle individual files
            for file_attr in ["ind_file1", "ind_file2", "ind_file3"]:
                file_field = getattr(mf, file_attr, None)
                if file_field:
                    url = request.build_absolute_uri(settings.MEDIA_URL + file_field.name)
                    file_urls.append(url)

        return JsonResponse({
            "meeting": meeting_data,
            "audio_files": file_urls,
            "transcript_files": transcript_file_urls,
            "transcript": transcript_text,
            "gemini": get_meeting_summary_and_tasks(meeting_data, transcript_text, transcript_file_urls)

        })

    except Meeting.DoesNotExist:
        return JsonResponse({"error": "Meeting not found"}, status=404)

import io
import os
import json
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.conf import settings
from django.core.files.base import ContentFile
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import simpleSplit

from .models import Meeting, MeetingFile, Employee, Task

@csrf_exempt
def approve_summary(request, meeting_id):
    if request.method == "POST":
        data = json.loads(request.body)
        summary = data.get("summary", [])
        tasks = data.get("tasks", {})

        # ✅ Fetch meeting first
        try:
            meeting = Meeting.objects.get(meeting_id=meeting_id)
        except Meeting.DoesNotExist:
            return JsonResponse({"error": "Meeting not found"}, status=404)

        # ✅ Resolve mic employees
        mic_employees = []
        for mic in [meeting.meeting_mic1, meeting.meeting_mic2, meeting.meeting_mic3]:
            if mic:
                try:
                    emp = Employee.objects.get(employee_id=mic)
                    mic_employees.append(emp.employee_name)
                except Employee.DoesNotExist:
                    mic_employees.append(f"Unknown (ID {mic})")
        # ✅ Resolve departments
        department_names = []
        if meeting.meeting_department:
            dept_ids = [d.strip() for d in meeting.meeting_department.split(",") if d.strip()]
            for did in dept_ids:
                try:
                    dept = Department.objects.get(department_id=did)
                    department_names.append(dept.department_name)
                except Department.DoesNotExist:
                    department_names.append(f"Unknown (ID {did})")

        # ✅ Resolve participants
            participants = []
            if meeting.meeting_participant:
                participant_ids = [p.strip() for p in meeting.meeting_participant.split(",") if p.strip()]
                for pid in participant_ids:
                    try:
                        emp = Employee.objects.get(employee_id=pid)
                        participants.append(emp.employee_name)
                    except Employee.DoesNotExist:
                        participants.append(f"Unknown (ID {pid})")


        # ✅ Generate PDF in memory
        buffer = io.BytesIO()
        p = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter

        # Define margins
        left_margin = 80
        right_margin = 80
        top_margin = 60

        p.setFont("Helvetica-Bold", 18)
        p.setFillColorRGB(0, 0, 1)  # Blue
        p.drawString(left_margin, height - top_margin, f"Meeting Summary of {meeting.meeting_title}")
        y = height - top_margin - 30

        p.setFont("Helvetica", 12)
        p.setFillColorRGB(0, 0, 0)  # Black
        # Meeting details
        details = [
        f"Title: {meeting.meeting_title}",
        f"Date: {meeting.meeting_date} {meeting.meeting_time}",
        f"Location: {meeting.meeting_location}",
        f"Department(s): {', '.join(department_names)}",
        f"Participants: {', '.join(participants)}",
        "Mic 1, Mic 2, Mic 3: " + ", ".join(mic_employees),
            ]

        for d in details:
            lines = simpleSplit(d, "Helvetica", 12, width - left_margin - right_margin)  # ✅ wrap
            for line in lines:
                p.drawString(left_margin, y, line)
                y -= 18
        y -= 15


        # Summary section
        p.setFont("Helvetica-Bold", 14)
        p.setFillColorRGB(0, 0, 1)  # Blue
        p.drawString(left_margin, y, "Summary:")
        y -= 20

        # Summary points (normal text)
        p.setFont("Helvetica", 11)
        p.setFillColorRGB(0, 0, 0)
        for point in summary:
            p.drawString(left_margin + 20, y, f"- {point}")
            y -= 15

        # Tasks section
        y -= 20
        p.setFont("Helvetica-Bold", 14)
        p.setFillColorRGB(0, 0, 1)  # Blue
        p.drawString(left_margin, y, "Tasks:")
        y -= 20

        # Task details
        p.setFont("Helvetica", 11)
        p.setFillColorRGB(0, 0, 0)
        for assignee, task_list in tasks.items():
            p.drawString(left_margin + 10, y, f"{assignee}:")
            y -= 18
            for i, task in enumerate(task_list, 1):  # ✅ add numbering
                task_header = f"{i}) Title: {task['task_title']}"
                lines = simpleSplit(task_header, "Helvetica", 11, width - left_margin - right_margin)
                for line in lines:
                    p.drawString(left_margin + 25, y, line)
                    y -= 15

                # Content
                content_text = f"    Content: {task['task_content']}"
                lines = simpleSplit(content_text, "Helvetica", 11, width - left_margin - right_margin)
                for line in lines:
                    p.drawString(left_margin + 40, y, line)
                    y -= 15
                    

                # Urgency + deadline
                p.drawString(left_margin + 25, y, f"    Urgency: {task['urgent_level']}")
                y -= 15

                #deadline
                p.drawString(left_margin + 25, y, f"    Deadline: {task['deadline']}")
                y -= 25


        p.save()
        buffer.seek(0)

        # ✅ Save PDF into MeetingFile table
        meeting_file = MeetingFile.objects.filter(meeting=meeting).first()
        file_name = f"meeting_{meeting.meeting_title}_summary.pdf"

        if meeting_file:
            meeting_file.meeting_summary.save(file_name, ContentFile(buffer.getvalue()), save=True)
        else:
            meeting_file = MeetingFile.objects.create(
                meeting=meeting,
                meeting_summary=ContentFile(buffer.getvalue(), file_name),
            )

        # ✅ Also save PDF into local folder (MEDIA_ROOT/transcripts)
        transcript_dir = os.path.join(settings.MEDIA_ROOT, "transcripts")
        os.makedirs(transcript_dir, exist_ok=True)
        local_path = os.path.join(transcript_dir, file_name)
        with open(local_path, "wb") as f:
            f.write(buffer.getvalue())

        # ✅ Save tasks into DB
        for assignee, task_list in tasks.items():
            try:
                participant = Employee.objects.get(employee_name=assignee)  # match by name
            except Employee.DoesNotExist:
                continue  # skip invalid assignee names

            for task in task_list:
                Task.objects.create(
                    task_title=task["task_title"],
                    task_content=task["task_content"],
                    urgent_level=task["urgent_level"],
                    deadline=task["deadline"] if task["deadline"] else None,
                    status="Pending",
                    assignee_id=participant.employee_id,
                )

        return JsonResponse({
            "message": "Summary, PDF, and tasks saved successfully!",
            "pdf_path": local_path
        })

    return JsonResponse({"error": "Invalid request method"}, status=405)


@csrf_exempt
def complaint_upload(request):
    if request.method != "POST":
        return JsonResponse({"error": "Invalid method"}, status=405)

    try:
        complaint_audio = request.FILES.get("complaint_audio")
        if not complaint_audio:
            return JsonResponse({"error": "Audio file required"}, status=400)

        # Step 1: Save the complaint first (audio only)
        complaint = Complaint.objects.create(
            complaint_date=request.POST.get("complaint_date"),
            employee_id=request.POST.get("employee_id"),
            customer_name=request.POST.get("customer_name"),
            customer_contact=request.POST.get("customer_contact"),
            complaint_audio=complaint_audio,
            status="In Progress"
        )

        # Step 2: Transcribe using the saved file path
        abs_path = complaint.complaint_audio.path
        transcript = azure_transcribe(abs_path)

        # Step 3: Save the transcript into DB
        complaint.complaint_transcript = transcript
        complaint.save()

        return JsonResponse({
            "message": "Complaint uploaded and transcribed",
            "complaint_id": complaint.pk,
            "audio_url": request.build_absolute_uri(complaint.complaint_audio.url),
            "transcript": transcript
        })

    except Exception as e:
        print("ERROR:", e)
        return JsonResponse({"error": str(e)}, status=500)